import searcher.py as sc

# You should download this library because it might not be into your python modules
import docx

data = sc.data

def change_rep(i,ip):
  cambios = []

  fuerza = sc.get_Fuerza(data["items"][i])
  unidad = sc.get_Unidad(data["items"][i])
  dependencia = sc.get_Dependencia(data["items"][i])
  ciudad = sc.get_Ciudad(data["items"][i]) 
  vlan = sc.get_VLAN(data["items"][i]) 
  proveedor = sc.get_Proveedor(data["items"][i])
  segmento = sc.get_Segmento(data["items"][i]) 
  mascara = sc.get_Mascara(data["items"][i])  
  gateway = sc.get_Gateway(data["items"][i])  
  dns = sc.get_DNS(data["items"][i]) 
  dire = sc.get_Dir(data["items"][i])

  for x in range(0,len(fuerza)):
    if fuerza[x] != 'NF':
      fuerza[x] = str(str(ip) + '(' + fuerza[x] + ')')
      unidad[x] = str(str(ip) + '(' + unidad[x]+ ')')
      dependencia[x] = str(str(ip) + '(' + dependencia[x]+ ')')
      ciudad[x] = str(str(ip) + '(' + ciudad[x]+ ')')
      vlan[x] = str(str(ip) + '(' + str(vlan[x])+ ')')
      proveedor[x] = str(str(ip) + '(' + proveedor[x]+ ')')
      segmento[x] = str(str(ip) + '(' + segmento[x]+ ')')
      mascara[x] = str(str(ip) + '(' + mascara[x]+ ')')
      gateway[x] = str(str(ip) + '(' + gateway[x]+ ')')
      dns[x] = str(str(ip) + '(' + dns[x]+ ')')
      dire[x] = str(str(ip) + '(' + dire[x]+ ')')
        

  cambios.append(fuerza)
  cambios.append(unidad)
  cambios.append(dependencia)
  cambios.append(ciudad)
  cambios.append(vlan)
  cambios.append(proveedor)
  cambios.append(segmento)
  cambios.append(mascara)
  cambios.append(gateway)
  cambios.append(dns)
  cambios.append(dire)

  return cambios


def Report():
  rep = docx.Document()
  rep.add_heading('Reportes SIEM',0)


  for i in range(0,len(data["items"])):
    if data['items'][i]['averageAlertRiskScore'] > 70:

      rep.add_heading("Amenaza "+str(i+1)+"\n", level=1)
      table = rep.add_table(rows = 1, cols = 2)
      table.style = 'Light List Accent 2'
      heading_cells = table.rows[0].cells

      keys1 = list(data["items"][i].keys())
      values = list(data["items"][i].values())
      values1 = []
      alert = data["items"][i].get('alertMeta')
      ip = alert.get('SourceIp')
      ip = ip[0]
      lista = change_rep(i,ip)       
      

      for y in values:

        if type(y)== list:
          for q in y:
            #print(q)
            if q=='NF':
              y.remove(q)
                    
        if type(y)== list:
          for q in y:
            #print(q)
            if q=='NF':
              y.remove(q)

        y = str(y)
        if y=="[]":
          y = "-1"
        
        values1.append(y)
        
      heading_cells[0].text = keys1[0]
      heading_cells[1].text = values1[0]

#---------------------------------  LLAMAR FUNCIONES DE SEARCH COMPARE ---------------------------------



      for x in range(1,len(data["items"][i].keys())):
          if values1[x]!="-1":
              cells = table.add_row().cells
              cells[0].text = keys1[x]
              cells[1].text = values1[x]


  rep.save("Reporte_SIEM.docx")
Report()


#  sc.get_Segmento(data["items"][i]) 

#  sc.get_Mascara(data["items"][i]) 
 
#  sc.get_Gateway(data["items"][i]) 
 
#  sc.get_DNS(data["items"][i]) 

#  sc.get_Dir(data["items"][i]) 

#  sc.get_OBSERV(data["items"][i]) 
